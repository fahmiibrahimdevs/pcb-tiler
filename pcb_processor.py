import io
import math
import pymupdf as fitz  # PyMuPDF
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm as reportlab_mm

# A4 Constants in mm
A4_WIDTH_MM = 210.0
A4_HEIGHT_MM = 297.0

# Margin Narrow in mm (0.5 inch = 12.7 mm)
DEFAULT_MARGIN_MM = 12.7

class PCBProcessor:
    """
    Core Engine for PCB PDF Extraction, Dimension Detection,
    Grid Layout Tiling, Ultra-Sharp 600 DPI Binarization,
    Pair Patterning: TOP (7 spaces) BOT [Dashed Cut Line] TOP (7 spaces) BOT,
    Horizontal Auto-Centering (Center), Row Spacing, and Cut Line Generation.
    Supports Independent Copy Inputs with Fixed Pair Layout Structure.
    """

    @staticmethod
    def inspect_pdf(pdf_bytes: bytes) -> dict:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages_info = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            rect = page.rect
            page_w_mm = round(rect.width * 25.4 / 72.0, 2)
            page_h_mm = round(rect.height * 25.4 / 72.0, 2)

            matrix = fitz.Matrix(600 / 72.0, 600 / 72.0)
            pix = page.get_pixmap(matrix=matrix)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            
            gray = img.convert("L")
            bw = gray.point(lambda p: 255 if p < 220 else 0)
            bbox = bw.getbbox()

            if bbox:
                left, top, right, bottom = bbox
                crop_w_px = right - left
                crop_h_px = bottom - top
                crop_w_mm = round((crop_w_px / 600.0) * 25.4, 2)
                crop_h_mm = round((crop_h_px / 600.0) * 25.4, 2)
                cropped_img = img.crop(bbox)
            else:
                crop_w_mm = page_w_mm
                crop_h_mm = page_h_mm
                cropped_img = img

            cropped_gray = cropped_img.convert("L")
            cropped_bw = cropped_gray.point(lambda p: 0 if p < 220 else 255, mode="1")

            buffer = io.BytesIO()
            cropped_bw.save(buffer, format="PNG")
            img_png_bytes = buffer.getvalue()

            pages_info.append({
                "page": page_num + 1,
                "page_width_mm": page_w_mm,
                "page_height_mm": page_h_mm,
                "content_width_mm": crop_w_mm,
                "content_height_mm": crop_h_mm,
                "detected_width_mm": crop_w_mm,
                "detected_height_mm": crop_h_mm,
                "png_bytes": img_png_bytes
            })

        doc.close()
        return {
            "total_pages": len(pages_info),
            "pages": pages_info
        }

    @staticmethod
    def render_pcb_image(pdf_bytes: bytes, page_num: int = 1, crop_content: bool = True) -> tuple:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if page_num < 1 or page_num > len(doc):
            page_num = 1
        page = doc[page_num - 1]

        matrix = fitz.Matrix(600 / 72.0, 600 / 72.0)
        pix = page.get_pixmap(matrix=matrix)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        rect = page.rect
        w_mm = rect.width * 25.4 / 72.0
        h_mm = rect.height * 25.4 / 72.0

        if crop_content:
            gray = img.convert("L")
            bw = gray.point(lambda p: 255 if p < 220 else 0)
            bbox = bw.getbbox()
            if bbox:
                left, top, right, bottom = bbox
                w_mm = round(((right - left) / 600.0) * 25.4, 2)
                h_mm = round(((bottom - top) / 600.0) * 25.4, 2)
                img = img.crop(bbox)

        img_gray = img.convert("L")
        img_bw = img_gray.point(lambda p: 0 if p < 220 else 255, mode="1")

        doc.close()
        return img_bw, round(w_mm, 2), round(h_mm, 2)

    @staticmethod
    def _build_sequence(items: list, layout_mode: str, slots_visibility: list = None) -> list:
        sequence = []
        slot_idx = 0

        if layout_mode == "pair_top_bot" and len(items) >= 2:
            top_candidates = [it for it in items if "BOT" not in it.get("label", "").upper() and "BOTTOM" not in it.get("label", "").upper()]
            bot_candidates = [it for it in items if "BOT" in it.get("label", "").upper() or "BOTTOM" in it.get("label", "").upper()]

            if top_candidates and bot_candidates:
                top_item = top_candidates[0]
                bot_item = bot_candidates[0]
            else:
                top_item = items[0]
                bot_item = items[1]

            max_copies = max(top_item.get("copies", 1), bot_item.get("copies", 1))

            for i in range(max_copies):
                if i < top_item.get("copies", 1):
                    vis_a = True if slots_visibility is None or slot_idx >= len(slots_visibility) else bool(slots_visibility[slot_idx])
                    sequence.append({
                        "item": top_item,
                        "is_pair_start": True,
                        "is_pair_end": False,
                        "visible": vis_a,
                        "slot_index": slot_idx
                    })
                    slot_idx += 1

                if i < bot_item.get("copies", 1):
                    vis_b = True if slots_visibility is None or slot_idx >= len(slots_visibility) else bool(slots_visibility[slot_idx])
                    sequence.append({
                        "item": bot_item,
                        "is_pair_start": False,
                        "is_pair_end": True,
                        "visible": vis_b,
                        "slot_index": slot_idx
                    })
                    slot_idx += 1
        else:
            for item in items:
                for c in range(item.get("copies", 1)):
                    vis = True if slots_visibility is None or slot_idx >= len(slots_visibility) else bool(slots_visibility[slot_idx])
                    sequence.append({
                        "item": item,
                        "is_pair_start": False,
                        "is_pair_end": False,
                        "visible": vis,
                        "slot_index": slot_idx
                    })
                    slot_idx += 1
        
        return sequence

    @staticmethod
    def _group_sequence_into_rows(sequence: list, layout_mode: str, printable_w_mm: float, gap_spaces: int = 7) -> list:
        pair_gap_mm = (gap_spaces / 7.0) * 5.0   # Dynamic gap inside pair from spaces input (7 spaces = 5mm)
        tab_gap_mm = 15.0   # TAB gap between pairs/units

        rows = []
        current_row = []
        current_w = 0.0

        for idx, entry in enumerate(sequence):
            item = entry["item"]
            w_mm = item["width_mm"]
            h_mm = item["height_mm"]

            gap_before = 0.0
            if current_w > 0:
                if layout_mode == "pair_top_bot" and not entry["is_pair_start"]:
                    gap_before = pair_gap_mm
                else:
                    gap_before = tab_gap_mm

            # Check if item overflows current row width, wrap to new row if needed
            if current_w + gap_before + w_mm > printable_w_mm + 0.1:
                rows.append(current_row)
                current_row = []
                current_w = 0.0
                gap_before = 0.0

            entry_data = dict(entry)
            entry_data["gap_before"] = gap_before
            current_row.append(entry_data)
            current_w += (gap_before + w_mm)

        if current_row:
            rows.append(current_row)

        return rows

    @staticmethod
    def generate_docx(
        items: list,
        gap_spaces: int = 7,
        tab_spaces: int = 12,
        row_gap_mm: float = 14.0,
        margin_mm: float = DEFAULT_MARGIN_MM,
        layout_mode: str = "pair_top_bot",
        show_cut_lines: bool = True,
        auto_center: bool = True,
        slots_visibility: list = None
    ) -> bytes:
        doc = Document()
        
        section = doc.sections[0]
        section.page_width = Mm(A4_WIDTH_MM)
        section.page_height = Mm(A4_HEIGHT_MM)
        section.top_margin = Mm(margin_mm)
        section.bottom_margin = Mm(margin_mm)
        section.left_margin = Mm(margin_mm)
        section.right_margin = Mm(margin_mm)

        spaces_str = " " * gap_spaces
        half_tab = max(2, tab_spaces // 2)
        tab_left_str = " " * half_tab
        tab_right_str = " " * half_tab

        sequence = PCBProcessor._build_sequence(items, layout_mode, slots_visibility=slots_visibility)
        printable_w_mm = A4_WIDTH_MM - (2 * margin_mm)
        rows = PCBProcessor._group_sequence_into_rows(sequence, layout_mode, printable_w_mm, gap_spaces=gap_spaces)

        blank_white_img = Image.new("L", (10, 10), 255)
        blank_buf = io.BytesIO()
        blank_white_img.save(blank_buf, format="PNG")
        blank_bytes = blank_buf.getvalue()

        for row_idx, row in enumerate(rows):
            p = doc.add_paragraph()
            if auto_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            p.paragraph_format.space_before = Pt(0)
            
            if show_cut_lines and row_idx < len(rows) - 1:
                p.paragraph_format.space_after = Pt((row_gap_mm / 2.0) * 2.83465)
            else:
                p.paragraph_format.space_after = Pt(row_gap_mm * 2.83465) if row_idx < len(rows) - 1 else Pt(0)

            p.paragraph_format.line_spacing = 1.0

            for entry_idx, entry in enumerate(row):
                item = entry["item"]
                is_start = entry["is_pair_start"]
                visible = entry.get("visible", True)
                img = item["image"]
                w_mm = item["width_mm"]
                h_mm = item["height_mm"]

                if entry_idx > 0:
                    if layout_mode == "pair_top_bot" and not is_start:
                        p.add_run(spaces_str)
                    else:
                        if show_cut_lines:
                            run_v = p.add_run(tab_left_str + "┆" + tab_right_str)
                            run_v.font.bold = True
                            run_v.font.size = Pt(11)
                        else:
                            p.add_run(" " * tab_spaces)

                run_img = p.add_run()
                if visible:
                    img_buf = io.BytesIO()
                    img.save(img_buf, format="PNG")
                    img_buf.seek(0)
                    run_img.add_picture(img_buf, width=Mm(w_mm), height=Mm(h_mm))
                else:
                    run_img.add_picture(io.BytesIO(blank_bytes), width=Mm(w_mm), height=Mm(h_mm))

            if show_cut_lines and row_idx < len(rows) - 1:
                p_cut = doc.add_paragraph()
                if auto_center:
                    p_cut.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cut.paragraph_format.space_before = Pt(0)
                p_cut.paragraph_format.space_after = Pt((row_gap_mm / 2.0) * 2.83465)
                run_line = p_cut.add_run("✂ - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -")
                run_line.font.size = Pt(10)
                run_line.font.name = 'Courier New'
                run_line.font.bold = True

        output_buf = io.BytesIO()
        doc.save(output_buf)
        return output_buf.getvalue()

    @staticmethod
    def generate_pdf(
        items: list,
        gap_spaces: int = 7,
        tab_gap_mm: float = 15.0,
        row_gap_mm: float = 14.0,
        margin_mm: float = DEFAULT_MARGIN_MM,
        layout_mode: str = "pair_top_bot",
        show_cut_lines: bool = True,
        auto_center: bool = True,
        slots_visibility: list = None
    ) -> bytes:
        output_buf = io.BytesIO()
        c = canvas.Canvas(output_buf, pagesize=A4)
        
        sequence = PCBProcessor._build_sequence(items, layout_mode, slots_visibility=slots_visibility)
        printable_w_mm = A4_WIDTH_MM - (2 * margin_mm)
        rows = PCBProcessor._group_sequence_into_rows(sequence, layout_mode, printable_w_mm, gap_spaces=gap_spaces)

        row_heights = []
        for row in rows:
            max_h = max(entry["item"]["height_mm"] for entry in row)
            row_heights.append(max_h)

        curr_y_mm = A4_HEIGHT_MM - margin_mm

        for row_idx, row in enumerate(rows):
            row_h_mm = row_heights[row_idx]
            
            if curr_y_mm - row_h_mm < margin_mm - 0.1:
                c.showPage()
                curr_y_mm = A4_HEIGHT_MM - margin_mm

            total_row_w_mm = 0.0
            for idx, entry in enumerate(row):
                total_row_w_mm += (entry["gap_before"] + entry["item"]["width_mm"])

            if auto_center and total_row_w_mm < A4_WIDTH_MM:
                curr_x_mm = (A4_WIDTH_MM - total_row_w_mm) / 2.0
            else:
                curr_x_mm = margin_mm

            draw_y_mm = curr_y_mm - row_h_mm

            for entry_idx, entry in enumerate(row):
                item = entry["item"]
                is_start = entry["is_pair_start"]
                visible = entry.get("visible", True)
                img = item["image"]
                w_mm = item["width_mm"]
                h_mm = item["height_mm"]
                gap_before = entry["gap_before"]

                if show_cut_lines and entry_idx > 0 and (layout_mode != "pair_top_bot" or is_start):
                    cut_x = (curr_x_mm + (gap_before / 2.0)) * reportlab_mm
                    c.saveState()
                    c.setDash(4, 4)
                    c.setStrokeColor(colors.HexColor('#0f172a')) # Dark solid black/navy
                    c.setLineWidth(1.8) # Ultra bold 1.8pt line
                    c.line(cut_x, (draw_y_mm - 3) * reportlab_mm, cut_x, (curr_y_mm + 3) * reportlab_mm)
                    c.restoreState()

                curr_x_mm += gap_before

                if visible:
                    img_buf = io.BytesIO()
                    img.save(img_buf, format="PNG")
                    img_buf.seek(0)

                    from reportlab.lib.utils import ImageReader
                    rl_img = ImageReader(img_buf)

                    c.drawImage(
                        rl_img,
                        curr_x_mm * reportlab_mm,
                        draw_y_mm * reportlab_mm,
                        width=w_mm * reportlab_mm,
                        height=h_mm * reportlab_mm,
                        mask='auto'
                    )

                curr_x_mm += w_mm

            if show_cut_lines and row_idx < len(rows) - 1:
                c.saveState()
                c.setDash(4, 4)
                c.setStrokeColor(colors.HexColor('#0f172a')) # Dark solid black/navy
                c.setLineWidth(1.8) # Ultra bold 1.8pt line
                cut_y = (draw_y_mm - (row_gap_mm / 2.0)) * reportlab_mm
                c.line(margin_mm * reportlab_mm, cut_y, (A4_WIDTH_MM - margin_mm) * reportlab_mm, cut_y)
                c.restoreState()

            curr_y_mm -= (row_h_mm + row_gap_mm)

        c.save()
        return output_buf.getvalue()
